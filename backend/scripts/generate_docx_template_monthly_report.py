"""
Script de generación del template DOCX del Informe Técnico Mensual.

Feature 022 (T022) — variante DOCX editable del Informe Técnico Mensual
(Grupo de Alto Rendimiento). Ejecutar una sola vez para (re)generar el
archivo binario:

  cd backend && .venv/bin/python scripts/generate_docx_template_monthly_report.py

El template resultante se guarda en:
  templates/documents/docx/training_monthly_technical_report.docx

Usa python-docx para construir la estructura OOXML y docxtpl (jinja2) para
las variables/loops. El contexto que consume este template es exactamente
el que arma ``build_report_document_context()`` en
``app/services/training/reports.py`` — MISMA fuente de verdad que el
template PDF gemelo (``templates/documents/pdf/training_monthly_technical_report.html``).
Cualquier cambio de nombres de variables en ese builder debe reflejarse
aquí y allá a la vez.

Variables consumidas (ver ``DocumentTemplateSpec`` en
``app/services/notification/template_registry.py`` para el set mínimo
requerido, y el docstring de ``build_report_document_context`` para el
resto, opcional/degradable):

  club_name            str
  month_label          str  "Mayo 2026"
  season_year          str  "2026"
  is_draft             bool
  header               dict|None   {project_name, executing_entity,
                                     report_responsible, period_label}
  sections             list[dict]|None  [{key, title, text, is_missing}]
                        en el orden aprobado (Objetivo, Plan de
                        entrenamiento, Desarrollo de actividades,
                        Participación en competencia, Resultados
                        obtenidos, Conclusiones).
  missing_sections     list[str]|None
  narrative_blocks     dict[str, dict]|None  (fallback legacy, FR-012)
  project_profile      dict|None  (fallback legacy, FR-012)
  metrics_snapshot     dict|None
  athlete_names        dict[str, str]|None  {str(id): "Nombre Apellido"}
  session_detail       dict|None  {rows, is_empty, placeholder}
  attendance_table     dict|None  {rows, is_empty, placeholder}
  competition_results  list[dict]|None  (raw, CompetitionResultItem)
  competition_groups   list[dict]|None  [{event_id, event_name,
                        event_date, series_kind, awards_points, results}]
  has_competition_results bool|None
  conjoint_sessions    list[dict]|None
  photos               list[dict]|None  [{data_uri, session_date,
                        caption, section}] — a nivel de DOCX el renderer
                        (T023) debe inyectar, además, un ``photo_image``
                        (``docxtpl.InlineImage``) por foto para que el
                        placeholder abajo se resuelva a una imagen real.
                        Este script de autoría NO hace ese binding
                        (se hace en tiempo de render, no de plantilla).

PRIVACIDAD (FR de la feature 022 y principios del proyecto): documento de
distribución restringida — contiene nombres reales de menores en las
tablas de asistencia y de resultados de competencia. Solo coach/admin.

COMPATIBILIDAD (FR-012): todas las claves "nuevas" (``header``,
``sections``, ``missing_sections``, ``session_detail``, ``attendance_table``,
``competition_groups``, ``has_competition_results``) se leen con
``is defined`` / ``or {}`` / ``or []`` para nunca reventar el render ante
un reporte generado antes de esta feature. Las ausencias se muestran como
"Pendiente de completar" (narrativa) o "Pendiente — regenerar informe"
(tablas), igual que en el template PDF gemelo.
"""

from __future__ import annotations

from pathlib import Path

OUTPUT_PATH = (
    Path(__file__).parents[1]
    / "templates"
    / "documents"
    / "docx"
    / "training_monthly_technical_report.docx"
)


def build_template() -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # -----------------------------------------------------------------------
    # Estilos globales
    # -----------------------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def add_note(text: str, size: int = 8) -> None:
        """Párrafo de aviso/pie discreto (gris, itálica, tamaño reducido)."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # -----------------------------------------------------------------------
    # Portada / encabezado institucional
    # -----------------------------------------------------------------------
    heading = doc.add_heading("", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Informe Técnico Mensual de Entrenamiento")
    run.font.color.rgb = RGBColor(0x2D, 0x50, 0x16)  # verde bosque

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Grupo de Alto Rendimiento — {{ club_name }}").bold = True
    sub.runs[0].font.size = Pt(12)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_p.add_run(
        "Período: {{ (header.period_label if header is defined and header "
        "else none) or month_label }} — Temporada {{ season_year }}"
    ).italic = True

    doc.add_paragraph()

    doc.add_heading("Datos del proyecto", level=2)
    header_table = doc.add_table(rows=0, cols=2)
    header_table.style = "Table Grid"
    header_rows = [
        ("Club", "{{ club_name }}"),
        (
            "Nombre del proyecto",
            "{{ (header.project_name if header is defined and header else none) "
            "or (project_profile.get('project_name') if project_profile else none) "
            "or '—' }}",
        ),
        (
            "Entidad ejecutora",
            "{{ (header.executing_entity if header is defined and header else none) "
            "or (project_profile.get('executing_entity') if project_profile else none) "
            "or '—' }}",
        ),
        (
            "Responsable del informe",
            "{{ (header.report_responsible if header is defined and header else none) "
            "or (project_profile.get('report_responsible') if project_profile else none) "
            "or '—' }}",
        ),
        (
            "Período",
            "{{ (header.period_label if header is defined and header else none) "
            "or month_label }}",
        ),
        ("Temporada", "{{ season_year }}"),
    ]
    for label, value in header_rows:
        row_cells = header_table.add_row().cells
        row_cells[0].paragraphs[0].add_run(label).bold = True
        row_cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Banner BORRADOR (condicional, {%p if %}) — FR-008
    # -----------------------------------------------------------------------
    doc.add_paragraph("{%p if is_draft %}")
    draft_p = doc.add_paragraph()
    draft_run = draft_p.add_run(
        "BORRADOR — Este informe está pendiente de aprobación por parte del "
        "entrenador. Los textos pueden estar incompletos o sin revisar. "
        "Secciones pendientes: "
        "{{ (missing_sections | join(', ')) if missing_sections is defined "
        "and missing_sections else 'ninguna' }}."
    )
    draft_run.bold = True
    draft_run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    doc.add_paragraph("{%p endif %}")

    add_note(
        "Documento de distribución restringida. Contiene datos personales de "
        "menores de edad protegidos por la Ley 1581 de 2012 (Colombia). Uso "
        "exclusivo del equipo técnico y administrativo del club. Prohibida su "
        "divulgación o reproducción sin autorización."
    )

    # -----------------------------------------------------------------------
    # Helper jinja: sections_by_key + macro de texto narrativo (mismo truco
    # que el template PDF gemelo: build_report_document_context entrega
    # ``sections`` como lista [{key, title, text, is_missing}]; se indexa acá
    # para poder pedir el texto de cada sección aprobada por clave. Degrada a
    # ``narrative_blocks`` (formato legacy) si ``sections`` no viene en el
    # contexto — compatibilidad FR-012 para reportes pre-feature-022).
    # -----------------------------------------------------------------------
    # NOTA: se evitan deliberadamente los marcadores de recorte de espacios
    # de Jinja2 (``{%- ... -%}``) dentro de este bloque. docxtpl aplica un
    # parche de texto/XML propio (``patch_xml``) antes de delegar a Jinja2;
    # en la práctica, combinar esos marcadores de recorte con un párrafo de
    # texto plano inmediatamente anterior corrompe esa fase de parcheo
    # (el macro se fusiona con el párrafo previo y se pierde el tag de
    # apertura — reproducido y confirmado en pruebas manuales con esta
    # versión de docxtpl). Los tags planos ``{% %}`` no presentan el
    # problema; el único costo es algo de espacio en blanco adicional
    # (irrelevante en Word, se colapsa visualmente).
    doc.add_paragraph(
        "{% set sections_by_key = {} %}"
        "{% for _sec in (sections if sections is defined and sections else []) %}"
        "{% set _ = sections_by_key.update({_sec.key: _sec}) %}"
        "{% endfor %}"
        "{% macro sec_text(key) %}"
        "{% set _s = sections_by_key.get(key) %}"
        "{% if _s %}{{ _s.text }}"
        "{% else %}"
        "{% set _legacy = (narrative_blocks or {}).get(key) %}"
        "{% if _legacy and _legacy.get('final_text') %}{{ _legacy.final_text }}"
        "{% else %}Pendiente de completar{% endif %}"
        "{% endif %}"
        "{% endmacro %}"
    )

    approved_sections = [
        ("objetivo", "1. Objetivo del período"),
        ("plan_entrenamiento", "2. Plan de entrenamiento"),
        ("desarrollo", "3. Desarrollo de actividades"),
        ("competencia", "4. Participación en competencia"),
        ("resultados", "5. Resultados obtenidos"),
        ("conclusiones", "6. Conclusiones"),
    ]
    for key, title in approved_sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph("{{ sec_text('%s') }}" % key)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Detalle de sesiones — tabla con fila templada ({%tr for %}).
    #
    # IMPORTANTE (docxtpl): el tag ``{%tr for %}`` y su cierre
    # ``{%tr endfor %}`` NO deben compartir fila con la fila de contenido —
    # cada uno vive en su PROPIA fila (que ``patch_xml`` elimina por
    # completo al procesar el tag), con la(s) fila(s) de contenido en medio.
    # Ponerlos en celdas de la misma fila de datos (probado manualmente)
    # corrompe el parseo XML — mismo patrón que ``{%p for/endfor %}`` con
    # párrafos.
    # -----------------------------------------------------------------------
    doc.add_heading("Detalle de sesiones", level=2)
    session_table = doc.add_table(rows=4, cols=5)
    session_table.style = "Table Grid"
    hdr_cells = session_table.rows[0].cells
    for i, label in enumerate(["Fecha", "Hora", "Foco técnico", "Lugar", "Asistencia"]):
        hdr_cells[i].paragraphs[0].add_run(label).bold = True

    session_table.rows[1].cells[0].paragraphs[0].add_run(
        "{%tr for row in (session_detail.rows if session_detail is defined "
        "and session_detail else []) %}"
    )

    body_cells = session_table.rows[2].cells
    body_cells[0].paragraphs[0].add_run("{{ row.session_date or '—' }}")
    body_cells[1].paragraphs[0].add_run("{{ row.start_time or '—' }}")
    body_cells[2].paragraphs[0].add_run("{{ row.technical_focus or '—' }}")
    body_cells[3].paragraphs[0].add_run("{{ row.location or '—' }}")
    body_cells[4].paragraphs[0].add_run(
        "{{ row.present_count }}/{{ row.attendee_total }}"
    )

    session_table.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    doc.add_paragraph(
        "{%p if (session_detail.is_empty if session_detail is defined and "
        "session_detail else true) %}"
    )
    doc.add_paragraph(
        "{{ (session_detail.placeholder if session_detail is defined and "
        "session_detail else none) or 'Pendiente — regenerar informe' }}"
    )
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Asistencia y rúbrica por atleta — tabla con fila templada
    # -----------------------------------------------------------------------
    doc.add_heading("Asistencia y rúbrica por atleta", level=2)
    add_note(
        "Documento interno del club. Contiene datos personales de menores — "
        "no distribuir externamente."
    )
    attendance_table_doc = doc.add_table(rows=4, cols=8)
    attendance_table_doc.style = "Table Grid"
    hdr_cells = attendance_table_doc.rows[0].cells
    for i, label in enumerate(
        ["Atleta", "Pres.", "Aus.", "Justif.", "Tarde", "Les.", "% Asist.", "Rúbrica (E/A/T)"]
    ):
        hdr_cells[i].paragraphs[0].add_run(label).bold = True

    attendance_table_doc.rows[1].cells[0].paragraphs[0].add_run(
        "{%tr for row in (attendance_table.rows if attendance_table is "
        "defined and attendance_table else []) %}"
    )

    body_cells = attendance_table_doc.rows[2].cells
    body_cells[0].paragraphs[0].add_run(
        "{{ (athlete_names.get(row.athlete_id) if athlete_names else none) "
        "or ('Atleta ' ~ loop.index) }}"
    )
    body_cells[1].paragraphs[0].add_run("{{ row.count_present }}")
    body_cells[2].paragraphs[0].add_run("{{ row.count_absent }}")
    body_cells[3].paragraphs[0].add_run("{{ row.count_justified }}")
    body_cells[4].paragraphs[0].add_run("{{ row.count_late }}")
    body_cells[5].paragraphs[0].add_run("{{ row.count_injured }}")
    body_cells[6].paragraphs[0].add_run(
        "{{ '%.1f'|format(row.attendance_pct or 0) }}%"
    )
    body_cells[7].paragraphs[0].add_run(
        "{{ '%.1f'|format(row.avg_rubric_effort) if row.avg_rubric_effort is not none "
        "else '—' }} / "
        "{{ '%.1f'|format(row.avg_rubric_attitude) if row.avg_rubric_attitude is not none "
        "else '—' }} / "
        "{{ '%.1f'|format(row.avg_rubric_technique) if row.avg_rubric_technique is not none "
        "else '—' }}"
    )

    attendance_table_doc.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    doc.add_paragraph(
        "{%p if (attendance_table.is_empty if attendance_table is defined and "
        "attendance_table else true) %}"
    )
    doc.add_paragraph(
        "{{ (attendance_table.placeholder if attendance_table is defined and "
        "attendance_table else none) or 'Pendiente — regenerar informe' }}"
    )
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Resultados de competencia — agrupados por evento ({%tr for %} a nivel
    # de evento + {%p for %} anidado dentro de la celda para las filas de
    # atleta). Evita anidar {%tr%} dentro de {%tr%} (no soportado de forma
    # confiable por docxtpl); un loop de párrafo anidado dentro de la celda
    # sí lo está.
    # -----------------------------------------------------------------------
    doc.add_heading("Participación en competencia — resultados", level=2)
    add_note(
        "Documento interno del club. Contiene datos personales de menores — "
        "no distribuir externamente."
    )
    comp_table = doc.add_table(rows=4, cols=3)
    comp_table.style = "Table Grid"
    hdr_cells = comp_table.rows[0].cells
    for i, label in enumerate(["Evento", "Fecha", "Resultados (categoría — atleta — pos. — pts.)"]):
        hdr_cells[i].paragraphs[0].add_run(label).bold = True

    comp_table.rows[1].cells[0].paragraphs[0].add_run(
        "{%tr for group in (competition_groups if competition_groups is "
        "defined and competition_groups else []) %}"
    )

    body_cells = comp_table.rows[2].cells
    body_cells[0].paragraphs[0].add_run("{{ group.event_name or 'Evento sin nombre' }}")
    body_cells[1].paragraphs[0].add_run(
        "{{ group.event_date or '—' }} "
        "({{ 'otorga puntos' if group.awards_points else 'no otorga puntos' }})"
    )
    # Celda de resultados: loop anidado {%p for%} (párrafo por resultado) —
    # un nivel de anidación {%p%} dentro de una fila {%tr%} sí es seguro;
    # anidar {%tr%} dentro de {%tr%} no lo es.
    results_cell = body_cells[2]
    results_cell.paragraphs[0].add_run("{%p for r in group.results %}")
    row_p = results_cell.add_paragraph()
    row_p.add_run(
        "{{ r.category or 'Sin categoría' }} — {{ r.athlete_name }} — "
        "Pos. {{ r.position if r.position else '—' }} — "
        "Pts. {{ r.points if r.points else '—' }}"
    )
    end_p = results_cell.add_paragraph()
    end_p.add_run("{%p endfor %}")

    comp_table.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    doc.add_paragraph(
        "{%p if not (has_competition_results if has_competition_results is "
        "defined else (competition_results | length > 0 if competition_results "
        "is defined and competition_results else false)) %}"
    )
    doc.add_paragraph("Sin competencias registradas en el período.")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Actividades conjuntas y salidas — complementaria (paridad con PDF).
    # No se envuelve la tabla en {%p if%} (ese tag opera sobre párrafos, no
    # sobre tablas); en su lugar, igual que session_detail/attendance_table,
    # la tabla siempre se renderiza (vacía si no hay filas) y un párrafo
    # aparte cubre el caso sin datos.
    # -----------------------------------------------------------------------
    doc.add_heading("Actividades conjuntas y salidas", level=2)
    conjoint_table = doc.add_table(rows=4, cols=5)
    conjoint_table.style = "Table Grid"
    hdr_cells = conjoint_table.rows[0].cells
    for i, label in enumerate(["Fecha", "Tipo", "Foco técnico", "Lugar", "Duración"]):
        hdr_cells[i].paragraphs[0].add_run(label).bold = True

    conjoint_table.rows[1].cells[0].paragraphs[0].add_run(
        "{%tr for s in (conjoint_sessions if conjoint_sessions is defined "
        "and conjoint_sessions else []) %}"
    )

    body_cells = conjoint_table.rows[2].cells
    body_cells[0].paragraphs[0].add_run("{{ s.date }}")
    body_cells[1].paragraphs[0].add_run(
        "{{ 'Actividad conjunta' if s.kind == 'actividad_conjunta' else 'Salida' }}"
    )
    body_cells[2].paragraphs[0].add_run("{{ s.technical_focus or '—' }}")
    body_cells[3].paragraphs[0].add_run("{{ s.location or '—' }}")
    body_cells[4].paragraphs[0].add_run("{{ s.duration_min }} min")

    conjoint_table.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    doc.add_paragraph(
        "{%p if not (conjoint_sessions if conjoint_sessions is defined else []) %}"
    )
    doc.add_paragraph("Sin actividades conjuntas ni salidas registradas en el período.")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Registro fotográfico — placeholder de imagen (T023 hace el binding de
    # InlineImage real en tiempo de render; acá solo se reserva la posición).
    # -----------------------------------------------------------------------
    doc.add_heading("Registro fotográfico", level=2)
    add_note(
        "Imágenes con consentimiento informado — no distribuir externamente "
        "(Ley 1581/2012)."
    )
    doc.add_paragraph("{%p for photo in (photos if photos is defined and photos else []) %}")
    caption_p = doc.add_paragraph()
    caption_p.add_run(
        "{{ photo.section or 'Grupo de Alto Rendimiento' }} — "
        "{{ photo.session_date }}{% if photo.caption %} — {{ photo.caption }}{% endif %}"
    ).bold = True
    image_p = doc.add_paragraph()
    image_p.add_run("{{ photo_image }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph(
        "{%p if not (photos if photos is defined and photos else []) %}"
    )
    doc.add_paragraph("Sin fotografías registradas para el período.")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Pie de documento
    # -----------------------------------------------------------------------
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "Informe elaborado por el equipo técnico del Club Deportivo Trocha y "
        "Ruta · Valle del Cauca, Colombia · Temporada {{ season_year }}."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_note(
        "Documento de distribución restringida. Contiene datos personales de "
        "menores de edad conforme a la Ley 1581 de 2012 y el Decreto 1377 de "
        "2013 (Colombia). Prohibida su reproducción o distribución sin "
        "autorización del responsable del informe.",
        size=7,
    )

    # -----------------------------------------------------------------------
    # Guardar
    # -----------------------------------------------------------------------
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"✓ Template DOCX generado: {OUTPUT_PATH}")
    print(f"  Tamaño: {OUTPUT_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    build_template()
