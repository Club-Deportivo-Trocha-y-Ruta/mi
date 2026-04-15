"""
Script de generación del template DOCX de autorización médica.

Paso 9 del workflow-notifications.
Ejecutar una sola vez para generar el archivo binario:
  cd backend && .venv/bin/python scripts/generate_docx_template.py

El template resultante se guarda en:
  templates/documents/docx/medical_clearance.docx

Usa python-docx para crear la estructura y docxtpl para las variables Jinja2.
Las variables disponibles son:
  {{ athlete_first_name }}, {{ athlete_last_name }}, {{ birth_date }},
  {{ club_name }}, {{ season_year }}, {{ medical_conditions }} (list)
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

OUTPUT_PATH = Path(__file__).parents[1] / "templates" / "documents" / "docx" / "medical_clearance.docx"


def build_template() -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -----------------------------------------------------------------------
    # Estilos globales
    # -----------------------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -----------------------------------------------------------------------
    # Encabezado
    # -----------------------------------------------------------------------
    heading = doc.add_heading("", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("{{ club_name }}")
    run.font.color.rgb = RGBColor(0x2D, 0x50, 0x16)  # verde bosque

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Autorización Médica para Participación Deportiva").bold = True
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    season_p = doc.add_paragraph()
    season_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    season_p.add_run("Temporada {{ season_year }}").italic = True

    doc.add_paragraph()  # espacio

    # -----------------------------------------------------------------------
    # Datos del atleta
    # -----------------------------------------------------------------------
    doc.add_heading("Datos del Atleta", level=2)

    fields = [
        ("Nombre", "{{ athlete_first_name }} {{ athlete_last_name }}"),
        ("Fecha de nacimiento", "{{ birth_date }}"),
        ("Club", "{{ club_name }}"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Condiciones médicas (tabla con loop docxtpl)
    # -----------------------------------------------------------------------
    doc.add_heading("Condiciones Médicas Relevantes", level=2)

    doc.add_paragraph(
        "Las siguientes condiciones médicas han sido declaradas por el padre/madre/acudiente y "
        "deben ser conocidas por el equipo técnico del club:"
    )

    # Lista de condiciones médicas (loop simple para evitar problemas de parsing XML con docxtpl)
    doc.add_paragraph("{% for cond in medical_conditions %}")
    p = doc.add_paragraph("• {{ cond }}")
    p.paragraph_format.left_indent = Pt(20)
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Declaración y firma del padre/acudiente
    # -----------------------------------------------------------------------
    doc.add_heading("Declaración del Padre / Madre / Acudiente", level=2)
    doc.add_paragraph(
        "Yo, en calidad de padre/madre/acudiente legal del/la atleta mencionado/a, declaro que "
        "la información proporcionada es verídica y autorizo su participación en las actividades "
        "deportivas organizadas por {{ club_name }} durante la temporada {{ season_year }}."
    )

    doc.add_paragraph()
    doc.add_paragraph("Firma del padre/madre/acudiente: ___________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Nombre: _____________________________________________  C.C.: ______________")
    doc.add_paragraph()
    doc.add_paragraph("Fecha: ____________________")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Firma del médico
    # -----------------------------------------------------------------------
    doc.add_heading("Aval Médico", level=2)
    doc.add_paragraph(
        "El/la médico firmante certifica que ha evaluado al/la atleta y que no existen "
        "contraindicaciones para su participación en actividades deportivas de formación."
    )
    doc.add_paragraph()
    doc.add_paragraph("Firma del médico: ___________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Nombre: _____________________________________________  Reg. Médico: ______________")
    doc.add_paragraph()
    doc.add_paragraph("Fecha: ____________________")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Pie de confidencialidad
    # -----------------------------------------------------------------------
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer_p.add_run(
        "⚠ Documento confidencial — contiene datos de menor de edad protegidos por ley. "
        "Uso exclusivo del personal autorizado de {{ club_name }}."
    )
    run_footer.italic = True
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # -----------------------------------------------------------------------
    # Guardar
    # -----------------------------------------------------------------------
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"✓ Template DOCX generado: {OUTPUT_PATH}")
    print(f"  Tamaño: {OUTPUT_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    build_template()
