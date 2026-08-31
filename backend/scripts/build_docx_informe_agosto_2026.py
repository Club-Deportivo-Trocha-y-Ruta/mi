"""
Genera el DOCX del Informe Mensual de Agosto 2026 a partir del borrador Markdown.

Script de un solo uso (borrador para revisión del entrenador). NO usa el
template de la plataforma (`generate_docx_template_monthly_report.py`), que se
alimenta de la BD vía docxtpl; este arma el documento directamente desde el
Markdown editado a mano y embebe las fotografías del mes.

    cd backend && .venv/bin/python scripts/build_docx_informe_agosto_2026.py

Entrada  : docs/11-informe-tecnico-mensual/referencias/Informe-Mensual-Tecnico-Agosto-2026-BORRADOR.md
Imágenes : docs/11-informe-tecnico-mensual/referencias/imagenes-agosto-2026/
Salida   : docs/11-informe-tecnico-mensual/referencias/Informe-Mensual-Tecnico-Agosto-2026-BORRADOR.docx

PRIVACIDAD: contiene nombres y fotografías de menores. La carpeta de destino
está en .gitignore a propósito. No versionar la salida.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).parents[2]
BASE = ROOT / "docs" / "11-informe-tecnico-mensual" / "referencias"
MD_PATH = BASE / "Informe-Mensual-Tecnico-Agosto-2026-BORRADOR.md"
IMG_DIR = BASE / "imagenes-agosto-2026"
OUT_PATH = BASE / "Informe-Mensual-Tecnico-Agosto-2026-BORRADOR.docx"

AZUL = RGBColor(0x1F, 0x38, 0x64)
GRIS = RGBColor(0x59, 0x59, 0x59)

# Fotografías por bloque. El texto del placeholder en el Markdown se busca por
# subcadena; cada entrada es (archivo, pie de foto).
PHOTO_BLOCKS: dict[str, list[tuple[str, str]]] = {
    "el grupo en entrenamiento": [
        ("9bd18c0a-a0ff-4094-84aa-09064f5e09b2.jpg",
         "El grupo antes de salir a rodar, con deportistas, familias y cuerpo técnico."),
        ("IMG_2458.jpg",
         "Salida del 23 de agosto: los filos de Pichindé, con Cali al fondo."),
    ],
    "sesión virtual de fortalecimiento": [
        ("IMG_2389.jpg", "Sesión virtual del 11 de agosto: cada deportista entrena desde su casa."),
        ("IMG_2390.jpg", "Trabajo de fuerza con apoyo en silla y butaco, guiado por el entrenador."),
        ("IMG_2391.jpg", "Cierre de la sesión: el grupo sostuvo el encuentro en plena emergencia."),
    ],
    "sesiones de entrenamiento del mes": [
        ("IMG_2254.jpg",
         "4 de agosto — Afloje post carrera en Mulaló y conversatorio de conclusiones de Palmira."),
        ("IMG_2407.jpg",
         "13 de agosto — Primera salida tras el terremoto: Sweet Spot y descenso técnico en La Buitrera."),
        ("VIDEO_2411_descenso-13ago.jpg",
         "13 de agosto — Descenso técnico sobre los filos de La Buitrera."),
        ("VIDEO_2449_caballos-sendero-20ago.jpg",
         "20 de agosto — Encuentro con caballos en el sendero: el grupo redujo la marcha y pasó sin presionarlos."),
        ("IMG_2459.jpg", "23 de agosto — Subida en Zona 2 hacia Pichindé."),
        ("e03f78c0-66bb-43f8-a5a8-d30ca424d810.jpg",
         "23 de agosto — Parada en el río Pichindé: deportistas, entrenadores y familias."),
        ("IMG_2496.jpg",
         "30 de agosto — Bike park La Paz–Golondrinas, con riders de descenso de la zona."),
        ("IMG_2498.jpg", "30 de agosto — El grupo en el mirador del bike park."),
    ],
    "campamento La Cumbre": [
        ("IMG_2297.jpg", "7 de agosto — Campamento montado por los propios deportistas."),
        ("IMG_2278.jpg", "7 de agosto — Fogata: el punto de encuentro de la noche."),
        ("IMG_2283.jpg", "7 de agosto — Malvaviscos asados alrededor del fuego."),
        ("IMG_2315.jpg", "8 de agosto — Desayuno compartido en la terraza de la finca."),
        ("IMG_2311.jpg", "8 de agosto — Amanecer con neblina sobre el valle."),
        ("IMG_2319.jpg", "8 de agosto — Recorrido en bicicleta por los caminos de la vereda Agua Clara."),
        ("IMG_2334.jpg", "8 de agosto — Ascenso en grupo entre plantaciones de pino, con apoyo de seguridad."),
        ("IMG_2325.jpg", "8 de agosto — Taller de avistamiento con binoculares y guía ilustrada de especies."),
        ("VIDEO_2330_grupo-observando.jpg", "8 de agosto — Ruta “La Selva”: el grupo sigue un ave en el dosel."),
        ("VIDEO_2330_ave-en-arbol.jpg", "8 de agosto — Una de las especies observadas durante el recorrido."),
    ],
    "V Válida Copa Valle Palmira": [
        ("PALMIRA_sophia-vargas-troncos-dia-1.jpg",
         "1 de agosto — Gymkhana: paso de troncos en el circuito del Bosque Municipal."),
        ("PALMIRA_miguel-anaya-obstaculo-dia-1.jpg",
         "1 de agosto — Gymkhana: precisión sobre el obstáculo, donde el crono pesa menos que no fallar."),
        ("PALMIRA_samuel-ortiz-circuito-dia-2.jpg",
         "2 de agosto — XCO: tramo adoquinado del circuito de 3,4 km."),
        ("PALMIRA_sofia-gomez-sophia-vargas-dia-2.jpg",
         "2 de agosto — Antes de la salida del XCO."),
        ("PALMIRA_juan-diego-garcia-xco-dia-2.jpg",
         "2 de agosto — Categoría Élite: la prueba más larga del fin de semana."),
        ("PALMIRA_equipo-grupal-dia-2.jpg",
         "2 de agosto — El equipo del club al cierre de la válida."),
    ],
    "podios y grupo": [
        ("PALMIRA_podio-prejuvenil-a-fem-dia-1.jpg",
         "1 de agosto — Podio de gymkhana Prejuvenil A Femenina: primer lugar para el club."),
        ("PALMIRA_isabel-quinones-medalla-dia-1.jpg",
         "1 de agosto — Cero faltas y el mejor tiempo del grupo limpio."),
        ("PALMIRA_podio-infantil-a-fem-dia-1.jpg",
         "1 de agosto — Podio de gymkhana Infantil A Femenina: segundo lugar."),
        ("PALMIRA_podio-mariana-isabel-dia-2.jpg",
         "2 de agosto — Podio compartido en XCO Prejuvenil A Femenina."),
        ("PALMIRA_podio-samuel-dia-2.jpg",
         "2 de agosto — Segundo lugar en XCO Prejuvenil A, en el regreso tras la lesión."),
        ("PALMIRA_podio-miguel-dia-2.jpg",
         "2 de agosto — Tercer lugar en XCO Infantil A, la categoría más numerosa del campeonato."),
    ],
}


def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bottom_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "1F3864")
    borders.append(bottom)
    p_pr.append(borders)


def _add_hyperlink(paragraph, text: str, url: str, size: int) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    for tag, attrs in (
        ("w:color", {"w:val": "1155CC"}),
        ("w:u", {"w:val": "single"}),
        ("w:sz", {"w:val": str(size * 2)}),
    ):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        r_pr.append(el)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)


def _emit(paragraph, text: str, *, size: int, color, bold: bool, italic: bool) -> None:
    """Escribe `text` resolviendo los [enlaces](url) que contenga."""
    for token in re.split(r"(\[[^\]]+\]\([^)]+\))", text):
        if not token:
            continue
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            _add_hyperlink(paragraph, link.group(1), link.group(2), size)
            continue
        run = paragraph.add_run(token)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color


def _inline(paragraph, text: str, *, size: int = 11, color=None) -> None:
    """Escribe texto con **negrita**, *cursiva*, `código` y [enlaces](url).

    Los enlaces se resuelven también cuando están anidados dentro de un tramo
    en negrita o cursiva (p. ej. ``*ver [sitio](https://...)*``).
    """
    pattern = r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)"
    for token in re.split(pattern, text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            _emit(paragraph, token[2:-2], size=size, color=color, bold=True, italic=False)
        elif token.startswith("*") and token.endswith("*"):
            _emit(paragraph, token[1:-1], size=size, color=color, bold=False, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            _emit(paragraph, token[1:-1], size=size, color=color, bold=False, italic=True)
        else:
            _emit(paragraph, token, size=size, color=color, bold=False, italic=False)


def add_photo_block(doc: Document, photos: list[tuple[str, str]]) -> None:
    """Inserta las fotos en una grilla de 2 columnas con su pie de foto."""
    usable = []
    for name, caption in photos:
        path = IMG_DIR / name
        if path.exists():
            usable.append((path, caption))
        else:
            print(f"  ! falta imagen: {name}")
    if not usable:
        return

    col_w = Cm(7.8)
    rows = (len(usable) + 1) // 2
    table = doc.add_table(rows=rows * 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (path, caption) in enumerate(usable):
        r, c = divmod(idx, 2)
        with Image.open(path) as im:
            w, h = im.size
        # Alto máximo 9 cm para que dos filas quepan cómodas en la página.
        width = col_w
        if h / w * 7.8 > 9.0:
            width = Cm(9.0 * w / h)

        cell = table.cell(r * 2, c)
        cell.width = col_w
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(path), width=width)

        cap_cell = table.cell(r * 2 + 1, c)
        cap_cell.width = col_w
        cap = cap_cell.paragraphs[0]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = GRIS
        cap.paragraph_format.space_after = Pt(10)

    # Celda sobrante cuando el número de fotos es impar.
    if len(usable) % 2:
        table.cell(rows * 2 - 2, 1).text = ""
        table.cell(rows * 2 - 1, 1).text = ""


def _keep_with_next(paragraph) -> None:
    """Evita que el párrafo se separe del contenido que le sigue."""
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def _no_row_split(row) -> None:
    """Impide que una fila se parta entre páginas."""
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    _no_row_split(header_row)
    # Repite el encabezado si la tabla llega a partirse de todos modos.
    tr_pr = header_row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))

    for i, text in enumerate(header):
        cell = header_row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "1F3864")
        _keep_with_next(para)

    for idx, row in enumerate(rows):
        table_row = table.add_row()
        _no_row_split(table_row)
        cells = table_row.cells
        for i, text in enumerate(row[: len(header)]):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            _inline(para, text, size=9)
            # Todas las filas menos la última se mantienen unidas a la siguiente,
            # de modo que la tabla completa salta de página en bloque.
            if idx < len(rows) - 1:
                _keep_with_next(para)

    doc.add_paragraph()


def parse_markdown(md: str) -> list[tuple[str, object]]:
    """Convierte el Markdown en una lista de (tipo, contenido)."""
    blocks: list[tuple[str, object]] = []
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            blocks.append(("rule", None))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append((f"h{level}", stripped.lstrip("#").strip()))
            i += 1
            continue

        if stripped.startswith("> "):
            blocks.append(("quote", stripped[2:]))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            cells = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in table_lines
                if not re.fullmatch(r"\|[\s:|-]+\|", row)
            ]
            if cells:
                blocks.append(("table", (cells[0], cells[1:])))
            continue

        if stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("bullets", items))
            continue

        blocks.append(("p", stripped))
        i += 1
    return blocks


def build() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)

    photo_pending: list[str] = []

    for kind, content in parse_markdown(md):
        if kind == "rule":
            para = doc.add_paragraph()
            _bottom_rule(para)

        elif kind == "h1":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(content).replace(" — BORRADOR", "").upper())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = AZUL
            _bottom_rule(para)

        elif kind in ("h2", "h3"):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14)
            run = para.add_run(str(content).upper())
            run.bold = True
            run.underline = True
            run.font.size = Pt(13 if kind == "h2" else 12)
            run.font.color.rgb = AZUL

        elif kind == "h4":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(10)
            run = para.add_run(str(content))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = AZUL

        elif kind == "quote":
            para = doc.add_paragraph()
            _inline(para, str(content), size=9, color=GRIS)

        elif kind == "bullets":
            for item in content:  # type: ignore[union-attr]
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(4)
                _inline(para, item)

        elif kind == "table":
            header, rows = content  # type: ignore[misc]
            add_table(doc, header, rows)

        elif kind == "p":
            text = str(content)

            if text.startswith("**[ ESPACIO PARA FOTOGRAF"):
                matched = next(
                    (key for key in PHOTO_BLOCKS if key.lower() in text.lower()), None
                )
                if matched:
                    add_photo_block(doc, PHOTO_BLOCKS[matched])
                else:
                    photo_pending.append(text)
                    para = doc.add_paragraph()
                    run = para.add_run(
                        re.sub(r"\*\*|\[|\]", "", text).strip()
                    )
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = GRIS
                continue

            if text.startswith("⚠️"):
                para = doc.add_paragraph()
                _inline(para, text, size=9, color=RGBColor(0xB0, 0x4A, 0x00))
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _inline(para, text)

    doc.save(OUT_PATH)
    print(f"OK -> {OUT_PATH}")
    if photo_pending:
        print("Bloques de foto sin imágenes asignadas:")
        for text in photo_pending:
            print("  -", re.sub(r"\*\*|\[|\]", "", text).strip())


if __name__ == "__main__":
    build()
